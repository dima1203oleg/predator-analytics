from __future__ import annotations

"""Multi-Source Document Processor.

Обробляє різні типи документів:
- Excel/CSV - структуровані дані
- PDF - текстові документи
- Images - OCR (Tesseract)
- Word (.docx) - документи
- Telegram - канали та повідомлення
- Websites - скрейпінг
- API - зовнішні джерела
"""

from datetime import UTC, datetime
import json
import logging
import os
from pathlib import Path
from typing import Any

logger = logging.getLogger("service.document_processor")


class DocumentProcessor:
    """Багатоджерельний процесор документів."""

    def __init__(self):
        self.temp_dir = os.getenv("TEMP_DIR", "/tmp/predator_docs")
        self.output_dir = os.getenv("OUTPUT_DIR", "/tmp/predator_processed")
        os.makedirs(self.temp_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)

        # Check available processors
        self._check_dependencies()

    def _check_dependencies(self):
        """Перевірка доступних бібліотек."""
        self.has_pypdf = False
        self.has_tesseract = False
        self.has_docx = False
        self.has_pandas = False
        self.has_telethon = False
        self.has_beautifulsoup = False

        try:
            import pypdf

            self.has_pypdf = True
        except ImportError:
            logger.warning("pypdf not available - PDF processing disabled")

        try:
            import pytesseract

            self.has_tesseract = True
        except ImportError:
            logger.warning("pytesseract not available - OCR disabled")

        try:
            import docx

            self.has_docx = True
        except ImportError:
            logger.warning("python-docx not available - Word processing disabled")

        try:
            import pandas as pd

            self.has_pandas = True
        except ImportError:
            logger.warning("pandas not available - Excel/CSV processing disabled")

        try:
            from telethon import TelegramClient

            self.has_telethon = True
        except ImportError:
            logger.warning("telethon not available - Telegram parsing disabled")

        try:
            from bs4 import BeautifulSoup

            self.has_beautifulsoup = True
        except ImportError:
            logger.warning("beautifulsoup4 not available - Web scraping limited")

    async def process_file(
        self, file_path: str, source_type: str, options: dict[str, Any] | None = None
    ) -> dict[str, Any]:
        """Обробка файлу відповідно до типу джерела."""
        options = options or {}
        result = {
            "source_type": source_type,
            "file_path": file_path,
            "processed_at": datetime.now(UTC).isoformat(),
            "status": "pending",
            "records": [],
            "metadata": {},
        }

        try:
            if source_type in ["excel", "csv"]:
                result = await self._process_excel_csv(file_path, options)
            elif source_type == "pdf":
                result = await self._process_pdf(file_path, options)
            elif source_type == "image":
                result = await self._process_image(file_path, options)
            elif source_type == "word":
                result = await self._process_word(file_path, options)
            else:
                result["status"] = "error"
                result["error"] = f"Невідомий тип джерела: {source_type}"

            result["status"] = "completed"

        except Exception as e:
            logger.exception(f"Помилка обробки {source_type}: {e}")
            result["status"] = "error"
            result["error"] = str(e)

        return result

    async def _process_excel_csv(self, file_path: str, options: dict[str, Any]) -> dict[str, Any]:
        """Обробка Excel та CSV файлів."""
        if not self.has_pandas:
            return {
                "status": "error",
                "error": "pandas не встановлено. Виконайте: pip install pandas openpyxl",
            }

        import pandas as pd

        # Визначення типу файлу
        ext = Path(file_path).suffix.lower()

        if ext == ".csv":
            df = pd.read_csv(file_path, encoding=options.get("encoding", "utf-8"))
        elif ext in [".xlsx", ".xls"]:
            sheet = options.get("sheet_name", 0)
            df = pd.read_excel(file_path, sheet_name=sheet)
        else:
            return {"status": "error", "error": f"Непідтримуване розширення: {ext}"}

        # Перетворення в записи
        records = df.to_dict(orient="records")

        return {
            "source_type": "excel" if ext != ".csv" else "csv",
            "file_path": file_path,
            "processed_at": datetime.now(UTC).isoformat(),
            "status": "completed",
            "records": records,
            "metadata": {
                "total_rows": len(df),
                "columns": list(df.columns),
                "file_size": os.path.getsize(file_path),
                "sheet_name": options.get("sheet_name", "default"),
            },
        }

    async def _process_pdf(self, file_path: str, options: dict[str, Any]) -> dict[str, Any]:
        """Обробка PDF документів."""
        if not self.has_pypdf:
            return {
                "status": "error",
                "error": "pypdf не встановлено. Виконайте: pip install pypdf",
            }

        from pypdf import PdfReader

        reader = PdfReader(file_path)
        pages = []
        full_text = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({"page_number": i + 1, "text": text, "char_count": len(text)})
            full_text.append(text)

        # Об'єднаний текст
        combined_text = "\n\n".join(full_text)

        return {
            "source_type": "pdf",
            "file_path": file_path,
            "processed_at": datetime.now(UTC).isoformat(),
            "status": "completed",
            "records": [{"full_text": combined_text, "pages": pages}],
            "metadata": {
                "total_pages": len(reader.pages),
                "total_chars": len(combined_text),
                "file_size": os.path.getsize(file_path),
                "pdf_info": dict(reader.metadata) if reader.metadata else {},
            },
        }

    async def _process_image(self, file_path: str, options: dict[str, Any]) -> dict[str, Any]:
        """OCR обробка зображень (Tesseract)."""
        if not self.has_tesseract:
            return {
                "status": "error",
                "error": "pytesseract не встановлено. Виконайте: pip install pytesseract pillow",
            }

        from PIL import Image
        import pytesseract

        # Налаштування мови
        lang = options.get("language", "ukr+eng")

        # Відкриття зображення
        image = Image.open(file_path)

        # OCR
        text = pytesseract.image_to_string(image, lang=lang)

        # Додатково: отримання даних з боксами
        data = pytesseract.image_to_data(image, lang=lang, output_type=pytesseract.Output.DICT)

        # Статистика
        words = [w for w in data["text"] if w.strip()]
        confidence_scores = [c for c in data["conf"] if c > 0]
        avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0

        return {
            "source_type": "image",
            "file_path": file_path,
            "processed_at": datetime.now(UTC).isoformat(),
            "status": "completed",
            "records": [{"extracted_text": text, "words": words}],
            "metadata": {
                "image_size": image.size,
                "mode": image.mode,
                "format": image.format,
                "word_count": len(words),
                "avg_confidence": round(avg_confidence, 2),
                "language": lang,
                "file_size": os.path.getsize(file_path),
            },
        }

    async def _process_word(self, file_path: str, options: dict[str, Any]) -> dict[str, Any]:
        """Обробка Word документів (.docx)."""
        if not self.has_docx:
            return {
                "status": "error",
                "error": "python-docx не встановлено. Виконайте: pip install python-docx",
            }

        from docx import Document

        document = Document(file_path)

        # Витягнення тексту
        paragraphs = []
        full_text = []

        for para in document.paragraphs:
            if para.text.strip():
                paragraphs.append(
                    {"text": para.text, "style": para.style.name if para.style else "Normal"}
                )
                full_text.append(para.text)

        # Таблиці
        tables = []
        for table in document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        combined_text = "\n\n".join(full_text)

        return {
            "source_type": "word",
            "file_path": file_path,
            "processed_at": datetime.now(UTC).isoformat(),
            "status": "completed",
            "records": [{"full_text": combined_text, "paragraphs": paragraphs, "tables": tables}],
            "metadata": {
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "char_count": len(combined_text),
                "file_size": os.path.getsize(file_path),
            },
        }

    async def process_url(
        self, url: str, source_type: str, options: dict[str, Any] | None = None
    ) -> dict[str, Any]:
        """Обробка URL джерел (Telegram, Website, API, RSS)."""
        options = options or {}

        if source_type == "telegram":
            return await self._process_telegram(url, options)
        if source_type == "website":
            return await self._process_website(url, options)
        if source_type == "api":
            return await self._process_api(url, options)
        if source_type == "rss":
            return await self._process_rss(url, options)
        return {"status": "error", "error": f"Невідомий тип URL джерела: {source_type}"}

    async def _process_telegram(self, channel: str, options: dict[str, Any]) -> dict[str, Any]:
        """Парсинг Telegram каналу."""
        # Для реального парсингу потрібен Telethon та API ключі
        # Тут реалізовано заглушку для демонстрації

        channel_name = channel.rsplit("/", maxsplit=1)[-1].replace("@", "")

        return {
            "source_type": "telegram",
            "channel": channel_name,
            "url": channel,
            "processed_at": datetime.now(UTC).isoformat(),
            "status": "pending_setup",
            "message": f"Канал {channel_name} додано до моніторингу. Налаштуйте TELEGRAM_API_ID та TELEGRAM_API_HASH для активації.",
            "records": [],
            "metadata": {"requires_api_keys": True, "channel_name": channel_name},
        }

    async def _process_website(self, url: str, options: dict[str, Any]) -> dict[str, Any]:
        """Скрейпінг веб-сайту."""
        import httpx

        try:
            async with httpx.AsyncClient(timeout=30) as client:
                response = await client.get(url, follow_redirects=True)
                response.raise_for_status()
                html = response.text

            # Базове витягнення тексту
            text = html
            title = ""
            description = ""

            if self.has_beautifulsoup:
                from bs4 import BeautifulSoup

                soup = BeautifulSoup(html, "html.parser")

                # Видалення скриптів та стилів
                for script in soup(["script", "style", "noscript"]):
                    script.decompose()

                text = soup.get_text(separator="\n", strip=True)
                title = soup.title.string if soup.title else ""

                # Meta description
                meta_desc = soup.find("meta", attrs={"name": "description"})
                if meta_desc:
                    description = meta_desc.get("content", "")

            return {
                "source_type": "website",
                "url": url,
                "processed_at": datetime.now(UTC).isoformat(),
                "status": "completed",
                "records": [
                    {
                        "url": url,
                        "title": title,
                        "description": description,
                        "text": text[:50000],  # Обмеження на 50К символів
                    }
                ],
                "metadata": {
                    "status_code": response.status_code,
                    "content_type": response.headers.get("content-type", ""),
                    "content_length": len(html),
                    "title": title,
                },
            }

        except Exception as e:
            return {"source_type": "website", "url": url, "status": "error", "error": str(e)}

    async def _process_api(self, url: str, options: dict[str, Any]) -> dict[str, Any]:
        """Отримання даних з API."""
        import httpx

        headers = {}
        if options.get("api_key"):
            headers["Authorization"] = f"Bearer {options['api_key']}"

        try:
            async with httpx.AsyncClient(timeout=30) as client:
                response = await client.get(url, headers=headers, follow_redirects=True)
                response.raise_for_status()

                # Спроба парсити як JSON
                try:
                    data = response.json()
                except json.JSONDecodeError:
                    data = {"raw_text": response.text}

            # Нормалізація даних
            records = data if isinstance(data, list) else [data]

            return {
                "source_type": "api",
                "url": url,
                "processed_at": datetime.now(UTC).isoformat(),
                "status": "completed",
                "records": records,
                "metadata": {
                    "status_code": response.status_code,
                    "content_type": response.headers.get("content-type", ""),
                    "record_count": len(records),
                },
            }

        except Exception as e:
            return {"source_type": "api", "url": url, "status": "error", "error": str(e)}

    async def _process_rss(self, url: str, options: dict[str, Any]) -> dict[str, Any]:
        """Парсинг RSS/Atom фідів."""
        import httpx

        try:
            async with httpx.AsyncClient(timeout=30) as client:
                response = await client.get(url, follow_redirects=True)
                response.raise_for_status()
                content = response.text

            # Базовий парсинг XML
            import xml.etree.ElementTree as ET

            root = ET.fromstring(content)

            items = []

            # RSS 2.0
            for item in root.findall(".//item"):
                items.append(
                    {
                        "title": item.findtext("title", ""),
                        "link": item.findtext("link", ""),
                        "description": item.findtext("description", ""),
                        "pubDate": item.findtext("pubDate", ""),
                    }
                )

            # Atom
            for entry in root.findall(".//{http://www.w3.org/2005/Atom}entry"):
                items.append(
                    {
                        "title": entry.findtext("{http://www.w3.org/2005/Atom}title", ""),
                        "link": entry.find("{http://www.w3.org/2005/Atom}link").get("href", "")
                        if entry.find("{http://www.w3.org/2005/Atom}link") is not None
                        else "",
                        "description": entry.findtext("{http://www.w3.org/2005/Atom}summary", ""),
                        "pubDate": entry.findtext("{http://www.w3.org/2005/Atom}updated", ""),
                    }
                )

            return {
                "source_type": "rss",
                "url": url,
                "processed_at": datetime.now(UTC).isoformat(),
                "status": "completed",
                "records": items,
                "metadata": {
                    "item_count": len(items),
                    "feed_type": "rss" if root.find(".//item") is not None else "atom",
                },
            }

        except Exception as e:
            return {"source_type": "rss", "url": url, "status": "error", "error": str(e)}


# Singleton
_processor: DocumentProcessor | None = None


def get_document_processor() -> DocumentProcessor:
    """Get or create the document processor singleton."""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
